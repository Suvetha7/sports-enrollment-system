from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "docs" / "Data_Driven_Sports_Enrollment_Final_Report.docx"
TEMPLATE_MEDIA = Path("/tmp/cse435_report/word/media")

STUDENT_NAME = "L SUVETHA"
REG_NO = "12222013"
COURSE = "Bachelor of Technology in Computer Science and Engineering"
FACULTY_NAME = "[Enter Faculty Name]"
FACULTY_DESIGNATION = "[Enter Designation]"
GITHUB_LINK = "[Add GitHub Repository Link]"
LINKEDIN_LINK = "[Add LinkedIn Profile Link]"
TOPIC = "Data-Driven Sports Enrollment and Management System"
SESSION = "January-April 2026"


def set_default_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.bold = True

    doc.sections[0].top_margin = Inches(1)
    doc.sections[0].bottom_margin = Inches(1)
    doc.sections[0].right_margin = Inches(1)
    doc.sections[0].left_margin = Inches(1.25)


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14 if level == 1 else 12)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if bold_prefix and text.startswith(bold_prefix):
        run1 = p.add_run(bold_prefix)
        run1.bold = True
        run1.font.name = "Times New Roman"
        run1.font.size = Pt(12)
        rest = text[len(bold_prefix):]
        run2 = p.add_run(rest)
        run2.font.name = "Times New Roman"
        run2.font.size = Pt(12)
        return
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and update the table in Word to refresh page numbers."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_separate, placeholder, fld_end])


def add_cover_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    if (TEMPLATE_MEDIA / "image1.jpeg").exists():
        p.add_run().add_picture(str(TEMPLATE_MEDIA / "image1.jpeg"), width=Inches(1.25))

    for line, bold, size in [
        ("Seminar Report", True, 16),
        ("On", False, 12),
        (TOPIC, True, 16),
        ("Submitted by", False, 12),
        (STUDENT_NAME, True, 14),
        (f"Registration No.: {REG_NO}", True, 12),
        (COURSE, True, 12),
        ("Under the Supervision of", False, 12),
        (FACULTY_NAME, True, 12),
        (FACULTY_DESIGNATION, False, 12),
        ("Lovely Professional University, Punjab", True, 12),
        ("April 2026", False, 12),
    ]:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing = 1.5
        run = para.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
    doc.add_page_break()


def add_declaration_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 1.5
    run = title.add_run("DECLARATION")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    declaration_paras = [
        f'I hereby declare that the seminar report titled "{TOPIC}" submitted in partial fulfillment of the requirements for the award of the degree of Bachelor of Technology in Computer Science and Engineering is a record of my own work carried out during the academic session {SESSION}.',
        "I further declare that this report has not been submitted, either in part or in full, to any other institution or university for the award of any degree or diploma.",
        "I confirm that the content of this report is original and prepared for academic evaluation. Any references used in the preparation of this report have been properly acknowledged in the references section.",
        "I take full responsibility for the authenticity, accuracy, and originality of the work presented in this report.",
    ]
    for para in declaration_paras:
        add_para(doc, para)

    add_para(doc, f"Name of the Student: {STUDENT_NAME}")
    add_para(doc, f"Registration Number: {REG_NO}")
    add_para(doc, f"Course: {COURSE}")
    add_para(doc, "Signature of the Student: __________________________")
    add_para(doc, "Date: __________________________")
    doc.add_page_break()


def add_certificate_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 1.5
    run = title.add_run("CERTIFICATE")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    certificate_text = (
        f"This is to certify that the seminar report entitled \"{TOPIC}\" submitted by "
        f"{STUDENT_NAME} bearing Registration Number {REG_NO} in partial fulfillment of the "
        "requirements for the award of the degree of Bachelor of Technology in Computer Science "
        "and Engineering is a bonafide record of work carried out under my supervision during "
        f"the academic session {SESSION}."
    )
    add_para(doc, certificate_text)
    add_para(doc, "The seminar work presented in this report is satisfactory in scope, quality, and academic relevance for evaluation.")
    add_para(doc, f"Faculty Name: {FACULTY_NAME}")
    add_para(doc, f"Designation: {FACULTY_DESIGNATION}")
    add_para(doc, "Signature: __________________________")
    add_para(doc, "Date: __________________________")
    doc.add_page_break()


def add_acknowledgement(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ACKNOWLEDGEMENT")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    paragraphs = [
        "I express my sincere gratitude to the faculty members of the Department of Computer Science and Engineering, Lovely Professional University, for providing the academic environment and support required to complete this seminar work successfully.",
        f"I am especially thankful to {FACULTY_NAME} for the guidance, motivation, and constructive feedback offered throughout the preparation of this report and the live project demonstration.",
        "I also acknowledge the contribution of online documentation, research articles, and technical references that helped in understanding database design, transactional processing, validation mechanisms, and dashboard-oriented analytics for software systems.",
        "Finally, I would like to thank my family and peers for their encouragement and support during the planning, implementation, testing, and documentation of this project.",
    ]
    for paragraph in paragraphs:
        add_para(doc, paragraph)
    doc.add_page_break()


def add_abstract(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ABSTRACT")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    abstract_paras = [
        "The Data-Driven Sports Enrollment and Management System is a web-based software solution designed to modernize sports program registration and administrative tracking. Traditional enrollment processes in sports training environments often rely on manual forms, spreadsheets, or fragmented digital systems, which lead to duplicate records, delayed seat updates, and a lack of real-time operational visibility. The proposed system addresses these limitations through a centralized architecture that integrates user registration, login authentication, program-wise seat management, transactional enrollment control, and analytical reporting.",
        "The implementation uses a modular Python backend, a relational database design, and an interactive web interface created with HTML, CSS, and JavaScript. The system validates student data before storage, protects passwords through hashing, enforces unique constraints on important identifiers, and updates program occupancy through atomic transactions. The dashboard presents real-time values such as total enrollments, active programs, occupancy rate, and category-wise participation trends. These features demonstrate how database consistency and software engineering principles can be applied to solve a practical management problem in the sports domain.",
        "Apart from operational automation, the project establishes a strong foundation for future analytics. The structured data captured by the system can support demand forecasting, anomaly detection, recommendation systems, and advanced visualization. Therefore, the project not only serves as a registration portal but also acts as a base for data-driven decision making in sports administration.",
    ]
    for paragraph in abstract_paras:
        add_para(doc, paragraph)
    doc.add_page_break()


def add_table_of_contents(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TABLE OF CONTENTS")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    add_toc(doc)
    doc.add_page_break()


def add_abbreviations(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LIST OF ABBREVIATIONS")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    abbreviations = [
        ("AI", "Artificial Intelligence"),
        ("API", "Application Programming Interface"),
        ("CSS", "Cascading Style Sheets"),
        ("DBMS", "Database Management System"),
        ("ER", "Entity Relationship"),
        ("HTML", "HyperText Markup Language"),
        ("HTTP", "HyperText Transfer Protocol"),
        ("JSON", "JavaScript Object Notation"),
        ("PBKDF2", "Password-Based Key Derivation Function 2"),
        ("SQL", "Structured Query Language"),
        ("UI", "User Interface"),
        ("UX", "User Experience"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.rows[0].cells[0].text = "Abbreviation"
    table.rows[0].cells[1].text = "Expansion"
    for short, full in abbreviations:
        cells = table.add_row().cells
        cells[0].text = short
        cells[1].text = full
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.5
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
    doc.add_page_break()


def add_project_info_table(doc: Document) -> None:
    add_heading(doc, "Project Snapshot", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    rows = [
        ("Project Title", TOPIC),
        ("Student Name", STUDENT_NAME),
        ("Registration Number", REG_NO),
        ("Course", COURSE),
        ("Implementation Stack", "Python, SQLite, HTML, CSS, JavaScript"),
        ("Project Nature", "Live working web-based management system with analytics"),
        ("Primary Use Case", "Sports program registration, seat tracking, and data analysis"),
    ]
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Details"
    for field, value in rows:
        cells = table.add_row().cells
        cells[0].text = field
        cells[1].text = value
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.5
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)


def add_chapter_1(doc: Document) -> None:
    add_heading(doc, "Chapter-1 Introduction", level=1)
    add_project_info_table(doc)
    add_heading(doc, "1.1 Introduction to the Topic", level=2)
    paras = [
        "Sports institutions, training academies, and university recreation centers are increasingly expected to operate with the same level of efficiency and data transparency as other digitally managed service environments. However, enrollment for sports programs is still frequently handled through paper forms, informal spreadsheets, or disconnected messaging channels. Such approaches are difficult to audit, vulnerable to duplication, and incapable of providing real-time occupancy status to stakeholders. The present project explores how a data-driven software platform can solve these operational challenges in a structured and scalable way.",
        "The chosen seminar topic, Data-Driven Sports Enrollment and Management System, focuses on using software engineering principles and database-backed workflows to automate a common administrative problem. The system centralizes registration, login, sports program listing, and enrollment confirmation while continuously reflecting actual seat availability. It also creates a data asset that can be analyzed for trends, utilization, and planning decisions. This makes the topic academically strong because it combines web engineering, database management systems, transaction control, input validation, and introductory analytics in one application.",
        "In a final-year Computer Science context, the topic is significant because it demonstrates how theoretical knowledge from subjects such as DBMS, software engineering, web technologies, and data analytics can be integrated to solve a domain-specific business problem. It is not limited to coding a webpage; instead, it involves understanding user requirements, designing tables and relationships, protecting data integrity, implementing backend services, and presenting results in an interactive way. The project therefore acts as both a seminar study and a working software prototype.",
    ]
    for para in paras:
        add_para(doc, para)
    add_heading(doc, "1.2 Background and Importance of the Topic", level=2)
    for para in [
        "The digital transformation of administrative systems has become essential in education, healthcare, logistics, and retail. Sports administration is undergoing a similar shift, especially as institutions try to optimize coaching resources, infrastructure usage, scheduling, and student engagement. Without digital support, administrators may not know how many seats are available in a program, whether a student has already enrolled, or which categories are attracting the highest demand. This results in poor planning, underutilized resources, and a weaker user experience.",
        "A data-driven enrollment system provides clear benefits. It standardizes the registration process, enforces validation at the point of data entry, maintains one authoritative source of truth, and enables reporting through structured queries. For users, the system improves transparency by showing available sports programs and their current occupancy. For administrators, it reduces repetitive manual work and supports evidence-based decision making. These advantages justify the relevance of the selected topic for modern software-enabled sports management.",
        "The topic is also important from a concurrency perspective. During peak registration periods, multiple students may attempt to enroll in the same program simultaneously. If updates are not transaction-safe, overbooking and inconsistent records may occur. Therefore, a reliable system must not only store data but also coordinate multi-step operations correctly. This project addresses that requirement through an atomic enrollment workflow that updates both enrollment records and seat counts in one unit of work.",
    ]:
        add_para(doc, para)
    add_heading(doc, "1.3 Problem Statement", level=2)
    add_para(doc, "Sports training organizations often depend on manual or semi-digital enrollment methods that create duplicate records, delayed capacity updates, weak visibility into utilization, and poor consistency under concurrent access. As a result, students may face uncertainty during registration, and administrators may make decisions using incomplete or outdated information. A centralized, transaction-aware, and analytics-ready system is required to improve the reliability and efficiency of sports enrollment operations.")
    add_heading(doc, "1.4 Objectives of the Seminar Work", level=2)
    objectives = [
        "To design and implement a centralized sports enrollment portal for student registration and login.",
        "To maintain accurate real-time seat availability for every sports program.",
        "To prevent duplicate user accounts and duplicate program enrollments through validation and database constraints.",
        "To use transactional processing so that critical operations remain consistent during concurrent access.",
        "To generate operational analytics that can support administrative decision making and future predictive extensions.",
        "To produce a live working project and a structured final report suitable for viva and academic evaluation.",
    ]
    for item in objectives:
        add_bullet(doc, item)
    add_heading(doc, "1.5 Scope of the Project", level=2)
    for para in [
        "The current implementation covers student account registration, authentication, sports program browsing, transactional enrollment, seat tracking, and analytical summaries. The system is intended for demonstration in a local deployment environment, which makes it easy to run during practical assessment or viva.",
        "The project does not currently include payment gateways, role-specific admin dashboards, email notifications, or large-scale production deployment features. These are deliberately treated as future scope so that the current implementation remains focused, achievable, and aligned with seminar expectations while still leaving room for extension.",
        "Because the architecture is modular, the same system can later be upgraded to support MySQL deployment, advanced dashboards, and predictive analytics modules without major redesign. Therefore, the scope is sufficient for academic demonstration while also showing long-term extensibility.",
    ]:
        add_para(doc, para)
    add_heading(doc, "1.6 Methodology Overview", level=2)
    for para in [
        "The methodology adopted in this project begins with identifying the operational problems present in manual sports registration systems. Based on those issues, the core entities such as users, sports programs, enrollments, and audit records were defined. A relational schema was designed to support unique constraints, referential integrity, and analytical querying.",
        "After the database design, the backend services were implemented in Python. These services were responsible for registration, authentication, data retrieval, enrollment processing, and dashboard generation. A web interface was then created using HTML, CSS, and JavaScript so that the project could be demonstrated as a complete live system rather than only a backend script. Finally, the workflow was tested using seeded data, unit tests, and API verification to ensure the system behaved as expected.",
    ]:
        add_para(doc, para)


def add_chapter_2(doc: Document) -> None:
    add_heading(doc, "Chapter-2 Literature Review", level=1)
    sections = {
        "2.1 Need for Centralized Digital Management": [
            "A review of contemporary administrative systems shows that centralized databases significantly reduce redundancy and improve traceability. When records are maintained in multiple spreadsheets or informal channels, inconsistencies become common because updates are not synchronized. Centralized systems eliminate this issue by ensuring that every transaction updates a single data source. This principle is directly relevant to sports registration, where user records, enrollment entries, and seat counts must remain synchronized.",
            "Research and industry practice both indicate that digital systems improve auditability and accountability. Instead of relying on verbal confirmation or manual counts, administrators can track every operation through stored records and timestamps. This is especially valuable in academic and training environments where program capacity and participation data may need to be reviewed later.",
        ],
        "2.2 Importance of Input Validation and Data Quality": [
            "Input validation plays a major role in reducing errors at the point of entry. Systems that accept unverified identifiers, inconsistent contact information, or repeated accounts produce unreliable datasets that later affect reporting and planning. Data quality is not a cosmetic concern; it directly influences operational decisions and system credibility.",
            "Scholarly discussions around data quality emphasize completeness, consistency, uniqueness, and validity. The current project addresses these factors through required fields, email pattern validation, unique student identifiers, and duplicate enrollment checks. This aligns the implementation with accepted data management principles discussed in database and information systems literature.",
        ],
        "2.3 Transaction Management in Multi-User Systems": [
            "A major concern in shared digital platforms is concurrency control. When multiple users interact with the same records at nearly the same time, inconsistent states may arise unless the system ensures atomicity and isolation. Classic database literature explains that ACID properties are essential for systems handling critical updates such as seat allocation, balance transfer, or reservation booking.",
            "The sports enrollment problem resembles a reservation system because seat availability must remain accurate. Therefore, an enrollment operation should not be treated as a simple insert statement. It must check the current occupancy, confirm availability, record the enrollment, update the count, and complete all steps successfully before the state becomes visible. This study incorporates that insight through transaction-safe logic.",
        ],
        "2.4 Web-Based Management Portals": [
            "Web-based portals are frequently selected in academic software projects because they are platform independent, easier to demonstrate, and familiar to end users. A browser-based system requires no complex installation on the client side and allows an integrated view of forms, dashboards, and reports.",
            "Literature on educational and service portals suggests that usability has a direct impact on adoption. Even technically correct systems may fail in practice if users find them difficult to understand. Hence, the present project combines backend rigor with a professional user interface so that the live demo reflects both software quality and usability considerations.",
        ],
        "2.5 Analytics as a Value-Adding Layer": [
            "Modern information systems increasingly treat operational data as a strategic asset. Once data is captured in structured form, analytical queries can reveal demand patterns, utilization rates, and process bottlenecks. The seminar topic is therefore not limited to transaction processing; it also explores how a registration system can evolve into a decision-support tool.",
            "This perspective is consistent with data analytics literature, which argues that meaningful insights often emerge from relatively simple aggregation, filtering, and trend analysis over well-designed databases. In the context of sports management, category-wise participation and occupancy metrics can help decide which programs should be expanded, restructured, or promoted.",
        ],
        "2.6 Research Gap and Relevance of the Present Work": [
            "Many studies discuss digital registration, transaction safety, or analytics individually, but practical academic prototypes often fail to connect these ideas into one coherent application. The present work bridges that gap by designing a live project that combines registration workflow, validation logic, seat tracking, and analytics in a single integrated system.",
            "This makes the project relevant for seminar evaluation because it is not purely theoretical and not merely a UI prototype. Instead, it demonstrates how literature-backed concepts can be converted into a working software artifact that solves a practical problem while also supporting future extension into advanced analytics.",
        ],
    }
    for heading, paragraphs in sections.items():
        add_heading(doc, heading, level=2)
        for para in paragraphs:
            add_para(doc, para)
    add_heading(doc, "2.7 Comparative Review of Existing Approaches", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = table.rows[0].cells
    headers[0].text = "Approach"
    headers[1].text = "Strengths"
    headers[2].text = "Limitations"
    headers[3].text = "Relevance to Present Work"
    entries = [
        ("Manual Forms", "Low setup effort", "High redundancy, no analytics", "Motivates automation"),
        ("Spreadsheet Tracking", "Easy to edit", "Weak concurrency control", "Shows need for database support"),
        ("Basic Registration Portals", "Online access", "Often limited insight generation", "Supports web-based design choice"),
        ("Full ERP Modules", "Comprehensive integration", "High complexity and cost", "Academic prototype offers focused alternative"),
    ]
    for row in entries:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.5
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)


def add_chapter_3(doc: Document) -> None:
    add_heading(doc, "Chapter-3 Conceptual Study / Seminar Work", level=1)
    sections = {
        "3.1 System Overview": [
            "The proposed system is a lightweight but complete web application that allows students to create accounts, authenticate themselves, browse sports programs, and confirm enrollments. The same system also exposes a dashboard that summarizes operational metrics. Although the current version is suitable for local execution during a live demonstration, its design follows common architectural separation between the user interface, backend logic, and database layer.",
            "The project is intentionally organized in modules so that each responsibility remains clear. The server layer receives requests and serves static files, the services layer contains business logic and validation, the database layer initializes connections and security helpers, and the frontend layer presents forms and analytics visually. This separation supports maintainability and makes the code easier to explain in viva.",
        ],
        "3.2 Functional Architecture": [
            "At the entry point, the Python server starts an HTTP service and exposes API endpoints for registration, login, enrollment, and dashboard data retrieval. The browser interface communicates with these endpoints using JavaScript fetch calls. Each request is converted into structured backend operations that validate data, execute database transactions, and produce JSON responses. The client then updates the page dynamically based on the returned values.",
            "This request-response cycle demonstrates a standard but important software architecture. It shows how frontend components, service logic, and storage interact while remaining loosely coupled. As a result, one can change the interface design without rewriting the enrollment algorithm, or migrate the database engine without changing the conceptual workflow.",
        ],
        "3.3 Database Design": [
            "The database contains four major tables: users, sports_programs, enrollments, and audit_log. The users table stores identity details such as name, email, student ID, phone number, and hashed password. The sports_programs table stores program-specific details including category, coach, capacity, current enrollment count, schedule, fee, skill level, and status.",
            "The enrollments table links users and programs. It records which student has joined which program, together with enrollment status, notes, and timestamp. A composite uniqueness rule ensures that the same user cannot be enrolled in the same program more than once. The audit_log table captures system and user events, which is useful for traceability and future monitoring. This schema is normalized enough for the project scope while still remaining simple to understand.",
        ],
        "3.4 Validation Strategy": [
            "Validation occurs at multiple levels in the system. At the interface level, the forms require necessary fields such as full name, email, student ID, and password. At the service level, additional checks verify email format, password length, and identifier completeness. At the database level, unique constraints reject duplicate email addresses and student IDs even if application logic is bypassed.",
            "This layered validation strategy is important because professional systems should not depend on a single line of defense. Frontend checks improve user experience, backend checks protect business logic, and database constraints preserve integrity at the persistence layer. The combination of these measures improves robustness and matches good software engineering practice.",
        ],
        "3.5 Transaction-Safe Enrollment Workflow": [
            "The most critical workflow in the project is program enrollment. The sequence begins by verifying that the student exists, then fetching the selected sports program, checking whether the program is open, checking whether seats are still available, and confirming that the student is not already enrolled. If all checks pass, the system inserts a new enrollment record and increments the enrolled_count field of the corresponding program.",
            "All of these steps execute inside one transaction. If any validation fails, the transaction is rolled back. This prevents inconsistent states such as a seat count being updated without a corresponding enrollment or duplicate enrollment records being created under race conditions. In viva, this is one of the strongest technical points of the project because it connects theory from DBMS to a visible real-world use case.",
        ],
        "3.6 Security and Authentication Considerations": [
            "Even though the project is a seminar prototype, basic security considerations were incorporated. Passwords are not stored in plain text; they are converted into PBKDF2-based hashes with random salts. During login, the entered password is hashed with the same salt and compared to the stored digest. This protects user credentials more effectively than direct string storage.",
            "The current version does not implement multi-factor authentication or complex authorization roles because those features are beyond the chosen scope. However, the code structure allows them to be added later. For the current objective, the inclusion of password hashing and validation logic is sufficient to demonstrate awareness of secure coding practices.",
        ],
        "3.7 Frontend and User Experience Design": [
            "The frontend is built using semantic HTML, professional CSS styling, and JavaScript-based asynchronous requests. The user interface contains a hero section, sports program cards, registration form, login form, enrollment form, category analytics, and recent activity feed. A refined color scheme and dashboard layout were used to ensure the project appears professional during live demonstration.",
            "This design decision is academically useful because the assessment requires not only code explanation but also a working live system. A clear interface helps the evaluator understand the system workflow quickly. It also demonstrates that software quality includes usability and presentation, not just backend correctness.",
        ],
        "3.8 Implementation Stack": [
            "The backend was developed in Python using standard library components, which keeps the project easy to run in a lab environment. SQLite was used as the local demonstration database because it does not require server installation, while a MySQL-compatible schema was separately prepared to align the system with enterprise-style relational deployment.",
            "The frontend uses HTML, CSS, and vanilla JavaScript. This combination keeps the application lightweight and understandable for academic presentation. It also makes the project repository easy to push to GitHub because the setup is minimal and the source files clearly reflect the role of each component.",
        ],
    }
    for heading, paragraphs in sections.items():
        add_heading(doc, heading, level=2)
        for para in paragraphs:
            add_para(doc, para)
    add_heading(doc, "3.9 Workflow Summary", level=2)
    workflow_steps = [
        "Student opens the portal and registers or logs in.",
        "The frontend sends the submitted data to the Python backend.",
        "The backend validates the data and checks business rules.",
        "For enrollment, the backend starts a transaction and checks seat availability and duplicates.",
        "The database updates the enrollment record and the current occupancy count.",
        "The dashboard refreshes and displays the latest system metrics.",
    ]
    for item in workflow_steps:
        add_bullet(doc, item)


def add_chapter_4(doc: Document) -> None:
    add_heading(doc, "Chapter-4 Results and Discussion", level=1)
    sections = {
        "4.1 Working Output of the System": [
            "The final system runs as a local web application accessible through a browser. On startup, the application initializes the database, creates the necessary tables, and seeds demonstration data such as sports programs, user accounts, and sample enrollments. The home page displays a polished dashboard interface with sections for metrics, sports programs, registration, authentication, and analytics.",
            "The successful startup and dashboard rendering confirm that the architecture is integrated correctly. The user can perform registration and login actions without directly interacting with database files or command-line scripts. This contributes to the project’s suitability for live demonstration because the evaluator can observe the end-to-end behavior directly through the browser interface.",
        ],
        "4.2 Registration and Authentication Results": [
            "The registration module accepts new user data and stores it only after validating the email, student ID, phone number, and password strength. When a duplicate email or student ID is entered, the system rejects the operation and returns a clear message. This confirms that both application-level and database-level protections are working correctly.",
            "The authentication module verifies existing users using hashed password comparison. Successful login returns the associated student details, while incorrect credentials result in an error response. These outcomes demonstrate that the system supports secure access and does not expose password data in plain text.",
        ],
        "4.3 Enrollment and Real-Time Seat Tracking": [
            "The enrollment module is the most functionally rich part of the project. When a valid user selects a program, the backend checks program status, current occupancy, and prior enrollment records before confirming the new entry. After a successful transaction, the remaining seat count decreases immediately and the dashboard reflects the new state.",
            "This behavior demonstrates real-time seat calculation and transactional consistency. In the live demonstration, it becomes clear that the system is not a static design mockup. Instead, it is an operational application that updates stored values and user-facing metrics in coordination.",
        ],
        "4.4 Analytics Dashboard Findings": [
            "The dashboard exposes summary metrics such as total users, total programs, total enrollments, aggregate capacity, filled seats, and occupancy rate. It also categorizes data by sports category and shows recent enrollment activity. These outputs reveal not only that the system stores data correctly, but also that it can convert operational data into meaningful management information.",
            "For example, category-wise counts make it possible to compare participation across football, badminton, swimming, athletics, basketball, and tennis. Such information can help administrators decide where to add resources, increase intake, or improve outreach. Even simple analytics therefore add significant value to the registration process.",
        ],
        "4.5 Testing and Verification": [
            "The project includes unit tests that validate critical business functions such as user registration, duplicate rejection, authentication, enrollment, and analytics summary generation. These tests passed successfully during verification. In addition, the server was executed manually and the dashboard API returned the expected data payload, confirming that the application works not only at the code level but also in live execution.",
            "Testing is important in seminar evaluation because it demonstrates that the project has been verified systematically rather than only by casual clicking. The inclusion of automated tests also increases credibility when explaining the code during viva.",
        ],
        "4.6 Advantages of the Proposed System": [
            "The system offers several clear benefits. It centralizes records, maintains a single source of truth, improves data quality through validation, prevents duplicate enrollments, and supports occupancy tracking in real time. It is easy to run, easy to demonstrate, and sufficiently modular for future upgrades. These features make it appropriate for academic submission and potentially useful as a starting prototype for real sports administration workflows.",
            "Another advantage is the balance between simplicity and technical depth. The project is not over-engineered, yet it still demonstrates practical concepts such as hashing, transactions, relational design, modular architecture, and analytics. This balance is ideal for a final-year presentation because it allows the student to explain the code confidently while still showing meaningful engineering decisions.",
        ],
        "4.7 Limitations of the Current Implementation": [
            "The current version is designed primarily for local demonstration and academic evaluation. It does not include deployment on a public server, advanced user roles, email integration, online payment support, or visual chart libraries. The use of SQLite in the live build also means it is optimized for portability rather than large-scale concurrent production usage.",
            "These limitations do not weaken the value of the project; instead, they define the boundary of the present work. In fact, clearly identifying limitations is an important academic practice because it shows awareness of how a prototype differs from a production-ready platform.",
        ],
        "4.8 Discussion and Insights Gained": [
            "A key insight from the project is that data integrity cannot be treated as an afterthought. Even a small management portal requires careful handling of unique identifiers, controlled updates, and validation logic. Another insight is that structured data immediately becomes more useful when supported by analytical queries. This transforms a simple operational system into a decision-support foundation.",
            "The project also reinforces the importance of modular coding and documentation. Separating the project into backend, database, frontend, and documentation components not only improves maintainability but also makes viva explanation easier. This confirms that good project organization is an academic as well as practical strength.",
        ],
    }
    for heading, paragraphs in sections.items():
        add_heading(doc, heading, level=2)
        for para in paragraphs:
            add_para(doc, para)
    add_heading(doc, "4.9 Sample Result Summary", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    header[0].text = "Metric"
    header[1].text = "Observed Behavior"
    header[2].text = "Interpretation"
    result_rows = [
        ("Registration", "New users created successfully", "System supports validated account creation"),
        ("Duplicate Check", "Repeated email or student ID rejected", "Data consistency is preserved"),
        ("Enrollment", "Seats updated after confirmation", "Real-time occupancy is maintained"),
        ("Dashboard", "Metrics and recent activities displayed", "Operational data is analytics-ready"),
        ("Testing", "Core unit tests passed", "Key logic is verified"),
    ]
    for row in result_rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.5
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)


def add_chapter_5(doc: Document) -> None:
    add_heading(doc, "Chapter-5 Conclusion and Future Scope", level=1)
    sections = {
        "5.1 Summary of the Seminar Work": [
            "This seminar work presented the design and implementation of a Data-Driven Sports Enrollment and Management System that addresses common issues in manual sports registration workflows. The project integrates registration, login, program management visibility, transactional enrollment, and analytics into a coherent web-based application.",
            "The work demonstrates how foundational Computer Science concepts can be translated into a practical solution. Instead of treating database design, validation, and analytics as separate topics, the project combines them into one end-to-end system that can be run and explained live during evaluation.",
        ],
        "5.2 Major Learning Outcomes": [
            "The project resulted in several important learning outcomes. First, it strengthened the understanding of relational database design and the role of keys, constraints, and normalization in preserving integrity. Second, it highlighted the practical importance of transactions and concurrency-aware updates in reservation-style workflows.",
            "Third, it improved software engineering skills related to modular architecture, code organization, and testing. Fourth, it showed how a user interface and backend services must cooperate to create a professional user experience. Finally, it demonstrated that even a moderate academic project can become more impactful when designed with future analytics in mind.",
        ],
        "5.3 Conclusions Drawn from the Study": [
            "The study confirms that a digital, centralized, and transaction-aware enrollment system can significantly improve reliability and visibility in sports management operations. Real-time seat tracking reduces confusion, unique constraints reduce duplication, and analytical summaries help convert operational data into managerial insight.",
            "The project also confirms that carefully scoped academic software can be both implementable and conceptually rich. By focusing on core workflows rather than attempting an oversized system, the final outcome remains presentable, functional, and academically defensible during viva.",
        ],
        "5.4 Future Scope": [
            "The current project can be extended in multiple directions. A natural next step is migration from local demo storage to a full MySQL or PostgreSQL deployment with role-based dashboards for administrators, coaches, and students. Another useful improvement would be a waitlist mechanism, automated email notifications, and downloadable reports for management.",
            "Beyond operational improvements, the system can be upgraded with predictive and analytical intelligence. Historical enrollment data can support demand forecasting, anomaly detection, and recommendation systems that suggest sports programs based on user preferences or past participation. Visual dashboards using chart libraries and business intelligence integration are also promising extensions.",
            "Because the present architecture already organizes data in a structured relational form, these future directions can be implemented incrementally. This makes the project not only a completed seminar topic but also a genuine platform for subsequent development or mini-research exploration.",
        ],
        "5.5 Professional Profile and Repository Details": [
            f"GitHub Project Repository Link: {GITHUB_LINK}",
            f"LinkedIn Profile Link: {LINKEDIN_LINK}",
            "The source code repository contains the backend, frontend, SQL scripts, tests, and documentation. Maintaining the project on GitHub supports version control, code presentation, and professional portfolio building.",
        ],
    }
    for heading, paragraphs in sections.items():
        add_heading(doc, heading, level=2)
        for para in paragraphs:
            add_para(doc, para)


def add_chapter_6(doc: Document) -> None:
    add_heading(doc, "Chapter-6 Implementation, Testing and Repository Readiness", level=1)
    sections = {
        "6.1 Source Code Organization": [
            "The source code of the project is arranged in a repository structure that supports maintainability, explainability, and version control. The root folder contains the execution entry point, documentation, SQL scripts, static assets, and unit tests. This organization is intentional because a final-year project should not appear as a collection of disconnected files; rather, it should reflect a disciplined development process that others can understand and extend.",
            "The src folder contains backend modules responsible for database setup, business logic, seed data, and server behavior. The static folder contains interface assets such as HTML, CSS, and JavaScript. The docs folder holds report-oriented documentation, while the tests folder captures functional verification logic. This structure supports GitHub publication because each concern is visibly separated, which is a desirable property for both academic review and professional presentation.",
        ],
        "6.2 Mapping of Files to Responsibilities": [
            "The run.py file acts as the entry point that starts the application server. The server.py module defines route handling and response behavior. The db.py module manages database creation, connection configuration, and password hashing utilities. The services.py module contains the core operational logic for registration, authentication, enrollment, and analytics. The seed_data.py module is responsible for inserting example records that make the live demonstration meaningful from the first run.",
            "On the client side, index.html defines the visible structure of the portal, style.css is responsible for the professional layout and color scheme, and app.js handles browser-side interaction with the backend APIs. This clean distribution of concerns is valuable in viva because the student can explain the software layer by layer rather than struggling with one oversized code file.",
        ],
        "6.3 Database Initialization and Seed Strategy": [
            "A useful implementation detail of this project is the automatic initialization of the database. On startup, the application checks whether the required data structures are already present; if not, it creates the tables and inserts a meaningful set of sample records. This approach makes the project highly portable because it reduces manual setup steps during live evaluation or GitHub cloning.",
            "The seeded data includes sample students, sports programs, and enrollments. This allows the dashboard to display realistic values immediately. It also provides an instant testing baseline because the evaluator can log in with a demo account or observe how enrollment counts are updated after a new operation. Seed data is therefore not only a convenience feature but also a practical support for demonstration and validation.",
        ],
        "6.4 Test Case Design": [
            "The unit tests were designed around high-risk workflows rather than trivial behaviors. The chosen test cases verify successful registration, rejection of duplicate registration, authentication of seeded users, successful enrollment in an available program, rejection of duplicate enrollment, and generation of analytics summaries. This set of tests covers the most important correctness concerns of the project.",
            "From an academic perspective, this approach is appropriate because it demonstrates understanding of verification priorities. Instead of only checking helper functions, the tests focus on the operations that affect system reliability and user-facing correctness. Such testing strategy strengthens the credibility of the project during evaluation.",
        ],
        "6.5 GitHub Readiness": [
            "The repository has been prepared in a GitHub-friendly form, which means the project includes a clear README, organized folders, test files, SQL scripts, and documentation that explain how to run the system. This is important because the assessment specifically expects the work to be presentable as a complete student project rather than a temporary local experiment.",
            "A project uploaded to GitHub serves multiple purposes. It supports version control, enables future improvements, acts as a portfolio artifact, and demonstrates professional discipline in source code management. In the context of placement preparation or academic review, this repository structure becomes an important complement to the live demonstration.",
        ],
        "6.6 Deployment and Demonstration Considerations": [
            "The current system is designed to run locally through a simple command-line start procedure. This makes it practical for lab presentations because there is minimal setup overhead. The evaluator can see the application launch, observe the dashboard, perform registration or enrollment, and inspect the corresponding behavior immediately.",
            "For broader deployment, the same codebase could be placed behind a production-ready WSGI or ASGI server, connected to a hosted relational database, and secured with role-based access control. These steps are not required for the current seminar, but the fact that they are achievable without redesign highlights the architectural strength of the present implementation.",
        ],
        "6.7 Professional Value of the Project": [
            "From a professional standpoint, the project reflects important industry-relevant skills: problem analysis, relational design, backend implementation, input validation, secure password handling, dynamic UI development, testing, and documentation. These are transferable competencies that extend beyond the sports domain.",
            "The repository also demonstrates that software quality includes how a project is presented. A system that is logically correct but poorly documented or badly organized may be difficult to evaluate. By contrast, this project is structured so that a reviewer can quickly understand the purpose, setup, architecture, and testing approach. That improves both academic defensibility and professional impression.",
        ],
    }
    for heading, paragraphs in sections.items():
        add_heading(doc, heading, level=2)
        for para in paragraphs:
            add_para(doc, para)


def add_references(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "References", level=1)
    references = [
        "Silberschatz, A., Korth, H. F., and Sudarshan, S. Database System Concepts. McGraw-Hill Education.",
        "Elmasri, R., and Navathe, S. B. Fundamentals of Database Systems. Pearson.",
        "Pressman, R. S., and Maxim, B. R. Software Engineering: A Practitioner’s Approach. McGraw-Hill.",
        "Connolly, T., and Begg, C. Database Systems: A Practical Approach to Design, Implementation, and Management. Pearson.",
        "Mozilla Developer Network. HTTP, HTML, CSS, and JavaScript documentation.",
        "Python Software Foundation. Python Standard Library Documentation.",
        "SQLite Documentation. Transaction Control and Database Features.",
        "Open Web Application Security Project (OWASP). Password Storage and Application Security Guidelines.",
    ]
    for ref in references:
        add_bullet(doc, ref)


def add_appendix(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "Appendix", level=1)
    add_heading(doc, "A.1 Important Files in the Project", level=2)
    files = [
        "run.py - application entry point",
        "src/server.py - server and API route handling",
        "src/services.py - business logic for registration, login, enrollment, and analytics",
        "src/db.py - database initialization and password hashing",
        "src/seed_data.py - demo data population",
        "static/index.html - page layout",
        "static/style.css - professional UI styling",
        "static/app.js - client-side interactivity",
        "sql/mysql_schema.sql - MySQL-compatible schema",
        "tests/test_services.py - unit tests for major workflows",
    ]
    for file_item in files:
        add_bullet(doc, file_item)
    add_heading(doc, "A.2 Viva Preparation Pointers", level=2)
    viva_points = [
        "Explain the problem statement before showing the code.",
        "Describe why transactions are needed in seat allocation workflows.",
        "Show how duplicate users and duplicate enrollments are prevented.",
        "Explain the reason for using SQLite in the live demo and MySQL in the report schema.",
        "Demonstrate the dashboard after one registration or enrollment action.",
    ]
    for point in viva_points:
        add_bullet(doc, point)
    add_heading(doc, "A.3 Likely Viva Questions with Brief Answer Hints", level=2)
    qa = [
        ("Why did you choose this topic?", "Because sports registration is a practical administrative problem where software, databases, and analytics can be meaningfully combined into one solution."),
        ("What is the main innovation in your project?", "The main innovation is combining transactional seat management with duplicate prevention and analytics in a live working portal."),
        ("How do you avoid duplicate entries?", "Unique constraints on email and student ID prevent duplicate users, while a user-program uniqueness rule prevents repeated enrollment in the same program."),
        ("Where is concurrency handled?", "Concurrency-sensitive enrollment is handled in the transaction-safe backend service where availability is checked and the seat count is updated in one atomic workflow."),
        ("Why is the project called data-driven?", "Because the system not only records transactions but also computes metrics, occupancy, and category-level insights from structured data."),
        ("Why did you use SQLite in the demo?", "SQLite makes local execution simple for assessment, while the included MySQL schema shows the project is designed for relational deployment beyond the demo environment."),
        ("What are the limitations of your system?", "The current version does not include payment integration, public deployment, admin roles, or advanced predictive analytics."),
        ("What future enhancements are possible?", "Future scope includes role-based dashboards, forecasting, recommendation systems, anomaly detection, waitlists, and richer visual analytics."),
    ]
    for question, hint in qa:
        add_para(doc, f"{question} {hint}")
    add_heading(doc, "A.4 Extended Discussion Notes for Viva", level=2)
    for para in [
        "While presenting the live project, it is useful to begin with the problem statement rather than the code. Explain that the system was created to address duplicate records, real-time seat management, and weak visibility in traditional sports registration processes. This allows the evaluator to understand the need for each technical component that follows.",
        "When showing the code, the recommended explanation order is entry point, server, services, database, frontend, and tests. This sequence mirrors the actual execution flow of the application and helps maintain clarity. If asked about technical depth, focus on the enrollment transaction and the rationale behind password hashing and unique constraints.",
        "During the report discussion, emphasize that the project is both operational and extensible. It already solves a real process problem, but it also generates structured data that can support future analytics. This dual perspective is one of the strongest conceptual aspects of the seminar topic and should be highlighted confidently.",
    ]:
        add_para(doc, para)


def build_report() -> Path:
    doc = Document()
    set_default_style(doc)
    for section in doc.sections:
        add_page_number(section)

    add_cover_page(doc)
    add_declaration_page(doc)
    add_certificate_page(doc)
    add_acknowledgement(doc)
    add_abstract(doc)
    add_table_of_contents(doc)
    add_abbreviations(doc)
    add_chapter_1(doc)
    doc.add_page_break()
    add_chapter_2(doc)
    doc.add_page_break()
    add_chapter_3(doc)
    doc.add_page_break()
    add_chapter_4(doc)
    doc.add_page_break()
    add_chapter_5(doc)
    doc.add_page_break()
    add_chapter_6(doc)
    add_references(doc)
    add_appendix(doc)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output_path = build_report()
    print(output_path)
